#!/usr/bin/env python
# coding=utf8


"""
====================================
 :mod:`argoslabs.office.wordeditor`
====================================
.. moduleauthor:: Arun Kumar <ak080495@gmail.com>
.. note:: ARGOS-LABS License

Description
===========
ARGOS-LABS RPA Word Editor module

"""
# Authors
# ===========
#
# * Arun Kumar ,
#
# Change Log
# --------
#
#  * [2022/07/24]
#     - icon
#     - input add
#     - read paragraph and table
#     - read table hide duplicate and empty
#     - update paragraph
#     - update paragraph path fix
#  * [2022/07/29]
#     - read update header and footer
#  * [2022/08/05]
#     - read plot
#  * [2022/08/08]
#     - add empty table row
################################################################################
import os
import sys
from alabs.common.util.vvargs import ModuleContext, func_log, \
    ArgsError, ArgsExit, get_icon_path
from io import StringIO
from docx import Document
import warnings


################################################################################
class Wordeditor(object):
    # ==========================================================================
    def __init__(self, argspec):
        self.argspec = argspec
        self.document = None

    # ==========================================================================
    def save_file(self,document, output_path, docx_file):
        location = docx_file
        if os.path.exists(location):
            base = os.path.basename(location)
            _a = os.path.splitext(base)[0]
            tem_location = os.path.join(output_path,
                                        "{0}"
                                        "({1})"
                                        "{2}".format(
                                            os.path.splitext(base)[0],
                                            str(1),
                                            os.path.splitext(base)[1]))
            location = self.create_dir(1, tem_location)
        _ = document.save(location)
        return location

    # ==========================================================================
    def create_dir(self,cf_ent, _location):
        if not os.path.exists(_location):
            return _location
        cf_ent += 1
        return self.create_dir(cf_ent,_location.replace("({0})".format(str(cf_ent - 1)),
                                                        "({0})".format(str(cf_ent))))

    # ==========================================================================
    def read_file(self):
        if self.argspec.docx_file:
            read_file = open(self.argspec.docx_file, 'rb')
            self.document = Document(read_file)
        return self.document

    # ==========================================================================
    def read_header(self):
        outst = StringIO()
        outst.write('p_id,paragraph')
        outst.write('\n')
        for s_data in self.document.sections:
            for h_p_id, h_p_data in enumerate(s_data.header.paragraphs):
                outst.write(f'{h_p_id},'
                            f'{h_p_data.text}'
                            )
                outst.write('\n')
        print(outst.getvalue(), end='')

    # ==========================================================================
    def read_header_run(self):
        outst = StringIO()
        outst.write('p_id,r_id,paragraph')
        outst.write('\n')
        for s_data in self.document.sections:
            for h_p_id, h_p_data in enumerate(s_data.header.paragraphs):
                for h_p_r_id, h_p_r_data in enumerate(h_p_data.runs):
                    outst.write(f'{h_p_id},'
                                f'{h_p_r_id},'
                                f'{h_p_r_data.text}')
                    outst.write('\n')
        print(outst.getvalue(), end='')

    # ==========================================================================
    def read_paragraphs(self):
        outst = StringIO()
        outst.write('p_id,paragraph')
        outst.write('\n')
        for p_id, p_data in enumerate(self.document.paragraphs):
            outst.write(f'{p_id},'
                        f'{p_data.text}'
                        )
            outst.write('\n')
        print(outst.getvalue(), end='')

    # ==========================================================================
    def read_paragraphs_nplot(self):
        for p_id, p_data in enumerate(self.document.paragraphs):
            p_data.text = f'({p_id})' \
                          f'\n' \
                          f'{p_data.text}' \
                          f'\n'
        output_path = os.path.dirname(self.argspec.docx_file)
        location = self.save_file(self.document, output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def read_paragraphs_plot(self):
        for p_id, p_data in enumerate(self.document.paragraphs):
            p_data.text = f'({p_id})' \
                          f'{p_data.text}'
        output_path = os.path.dirname(self.argspec.docx_file)
        location = self.save_file(self.document, output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def read_paragraphs_run(self):
        outst = StringIO()
        outst.write('p_id,r_id,paragraph')
        outst.write('\n')
        for p_id, p_data in enumerate(self.document.paragraphs):
            for r_id, r_data in enumerate(self.document.paragraphs[p_id].runs):
                outst.write(f'{p_id},'
                            f'{r_id},'
                            f'{r_data.text}'
                            )
                outst.write('\n')
        print(outst.getvalue(), end='')

    # ==========================================================================
    def read_paragraphs_run_nplot(self):
        for p_id, p_data in enumerate(self.document.paragraphs):
            for r_id, r_data in enumerate(self.document.paragraphs[p_id].runs):
                r_data.text = f'({p_id},' \
                              f'{r_id})' \
                              f'\n' \
                              f'{r_data.text}' \
                              f'\n'
        output_path = os.path.dirname(self.argspec.docx_file)
        location = self.save_file(self.document, output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def read_paragraphs_run_plot(self):
        for p_id, p_data in enumerate(self.document.paragraphs):
            for r_id, r_data in enumerate(self.document.paragraphs[p_id].runs):
                r_data.text = f'({p_id},' \
                              f'{r_id})' \
                              f'{r_data.text}'
        output_path = os.path.dirname(self.argspec.docx_file)
        location = self.save_file(self.document, output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    @staticmethod
    def read_table_run_nplot_next(j, r_id, c_id):
        for rp_id, k in enumerate(j.paragraphs):
            for rr_id, l in enumerate(k.runs):
                l.text = f'(' \
                         f'{r_id},' \
                         f'{c_id},' \
                         f'{rp_id},' \
                         f'{rr_id}' \
                         f')' \
                         f'\n' \
                         f'{l.text}' \
                         f'\n'

    # ==========================================================================
    def read_table_run_nplot(self):
        duplicate_hide=''
        for r_id, i in enumerate(self.document.tables[self.argspec.table_id].rows):
            for c_id, j in enumerate(i.cells):
                if j.text != duplicate_hide and j.text != '':
                    duplicate_hide = j.text
                    self.read_table_run_nplot_next(j, r_id, c_id)
        output_path = os.path.dirname(self.argspec.docx_file)
        location = self.save_file(self.document,
                                  output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    @staticmethod
    def read_table_run_plot_next(j, r_id, c_id):
        for rp_id, k in enumerate(j.paragraphs):
            for rr_id, l in enumerate(k.runs):
                l.text = f'(' \
                         f'{r_id},' \
                         f'{c_id},' \
                         f'{rp_id},' \
                         f'{rr_id}' \
                         f')' \
                         f'{l.text}'

    # ==========================================================================
    def read_table_run_plot(self):
        duplicate_hide=''
        for r_id, i in enumerate(self.document.tables[self.argspec.table_id].rows):
            for c_id, j in enumerate(i.cells):
                if j.text != duplicate_hide and j.text != '':
                    duplicate_hide = j.text
                    self.read_table_run_plot_next(j, r_id, c_id)
        output_path = os.path.dirname(self.argspec.docx_file)
        location = self.save_file(self.document,
                                  output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def read_table_nplot(self):
        duplicate_hide=''
        for r_id, i in enumerate(self.document.tables[self.argspec.table_id].rows):
            for c_id, j in enumerate(i.cells):
                if j.text != duplicate_hide and j.text != '':
                    duplicate_hide = j.text
                    j.text = f'(' \
                             f'{r_id},' \
                             f'{c_id}' \
                             f')' \
                             f'\n' \
                             f'{j.text}' \
                             f'\n'
        output_path = os.path.dirname(self.argspec.docx_file)
        location = self.save_file(self.document,
                                  output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def read_table_plot(self):
        duplicate_hide=''
        for r_id, i in enumerate(self.document.tables[self.argspec.table_id].rows):
            for c_id, j in enumerate(i.cells):
                if j.text != duplicate_hide and j.text != '':
                    duplicate_hide = j.text
                    j.text = f'(' \
                             f'{r_id},' \
                             f'{c_id}' \
                             f')' \
                             f'{j.text}'
        output_path = os.path.dirname(self.argspec.docx_file)
        location = self.save_file(self.document,
                                  output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    @staticmethod
    def read_table_run_next(j, outst, r_id, c_id):
        for rp_id, k in enumerate(j.paragraphs):
            for rr_id, l in enumerate(k.runs):
                outst.write(
                    f'{r_id},'
                    f'{c_id},'
                    f'{rp_id},'
                    f'{rr_id},'
                    f'{l.text}'
                )
                outst.write('\n')

    # ==========================================================================
    def read_table_run(self):
        outst = StringIO()
        outst.write('row_id,cell_id,p_id,r_id,value')
        duplicate_hide = ''
        for r_id, i in enumerate(self.document.tables[self.argspec.table_id].rows):
            for c_id, j in enumerate(i.cells):
                if j.text != duplicate_hide and j.text != '':
                    duplicate_hide = j.text
                    self.read_table_run_next(j, outst, r_id, c_id)
        print(outst.getvalue(), end='')

    # ==========================================================================
    def read_table(self):
        outst = StringIO()
        outst.write('row_id,cell_id,value')
        duplicate_hide = ''
        for r_id, i in enumerate(self.document.tables[self.argspec.table_id].rows):
            for c_id, j in enumerate(i.cells):
                if j.text != duplicate_hide and j.text != '':
                    duplicate_hide = j.text
                    outst.write(
                        f'{r_id},'
                        f'{c_id},'
                        f'{j.text}'
                    )
                    outst.write('\n')
        print(outst.getvalue(), end='')

    # ==========================================================================
    def read_footer_run(self):
        outst = StringIO()
        outst.write('p_id,r_id,paragraph')
        outst.write('\n')
        for s_data in self.document.sections:
            for f_p_id, f_p_data in enumerate(s_data.footer.paragraphs):
                for f_p_r_id, f_p_r_data in enumerate(f_p_data.runs):
                    outst.write(f'{f_p_id},'
                                f'{f_p_r_id},'
                                f'{f_p_r_data.text}'
                                )
                    outst.write('\n')
        print(outst.getvalue(), end='')

    # ==========================================================================
    def read_footer(self):
        outst = StringIO()
        outst.write('p_id,paragraph')
        outst.write('\n')
        for s_data in self.document.sections:
            for f_p_id, f_p_data in enumerate(s_data.footer.paragraphs):
                outst.write(f'{f_p_id},'
                            f'{f_p_data.text}'
                            )
                outst.write('\n')
        print(outst.getvalue(), end='')

    # ==========================================================================
    def write_header_run(self):
        for index, id in enumerate(self.argspec.paragrap_id):
            self.document.sections[0].header.paragraphs[id].runs[
                self.argspec.run_id[index]].text = \
                self.argspec.update_value[index]
        location = self.save_file(self.document,
                                  self.argspec.output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def write_header(self):
        for index, id in enumerate(self.argspec.paragrap_id):
            self.document.sections[0].header.paragraphs[id].text = \
                self.argspec.update_value[index]
        location = self.save_file(self.document,
                                  self.argspec.output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def write_paragraphs_run(self):
        for index, id in enumerate(self.argspec.paragrap_id):
            self.document.paragraphs[id].runs[self.argspec.run_id[index]].text = \
                self.argspec.update_value[index]
        location = self.save_file(self.document,
                                  self.argspec.output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def write_paragraphs(self):
        for index, id in enumerate(self.argspec.paragrap_id):
            self.document.paragraphs[id].text = self.argspec.update_value[index]
        location = self.save_file(self.document,
                                  self.argspec.output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def write_table_run(self):
        for index,r_id in enumerate(self.argspec.row_id):
            r_c_data = self.document.tables[int(self.argspec.table_id)]. \
                rows[r_id]. \
                cells[self.argspec.cell_id[index]]. \
                paragraphs[self.argspec.paragrap_id[index]]
            if len(r_c_data.runs) == 0:
                r_c_data.add_run(self.argspec.update_value[index])
            else:
                r_c_data.runs[self.argspec.run_id[index]].text = \
                    self.argspec.update_value[index]
        location = self.save_file(self.document,
                                  self.argspec.output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def write_table(self):
        for index,r_id in enumerate(self.argspec.row_id):
            r_c_data = self.document.tables[int(self.argspec.table_id)]. \
                rows[r_id].cells[self.argspec.cell_id[index]]
            r_c_data.text = self.argspec.update_value[index]
        location = self.save_file(self.document,
                                  self.argspec.output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def write_footer_run(self):
        for index, id in enumerate(self.argspec.paragrap_id):
            self.document.sections[0].footer.paragraphs[id].runs[
                self.argspec.run_id[index]].text = \
                self.argspec.update_value[index]
        location = self.save_file(self.document,
                                  self.argspec.output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def write_footer(self):
        for index, id in enumerate(self.argspec.paragrap_id):
            self.document.sections[0].footer.paragraphs[id].text = \
                self.argspec.update_value[index]
        location = self.save_file(self.document,
                                  self.argspec.output_path,
                                  self.argspec.docx_file)
        print(location, end='')

    # ==========================================================================
    def do(self, op):
        if op == 'Read Header':
            if self.argspec.run:
                self.read_header_run()
            else:
                self.read_header()
        elif op == 'Read Paragraphs':
            if self.argspec.run and self.argspec.nplot:
                self.read_paragraphs_run_nplot()
            elif self.argspec.run and self.argspec.plot:
                self.read_paragraphs_run_plot()
            elif self.argspec.nplot:
                self.read_paragraphs_nplot()
            elif self.argspec.plot:
                self.read_paragraphs_plot()
            elif self.argspec.run:
                self.read_paragraphs_run()
            else:
                self.read_paragraphs()
        elif op == 'Read Table':
            if self.argspec.run and self.argspec.nplot:
                self.read_table_run_nplot()
            elif self.argspec.run and self.argspec.plot:
                self.read_table_run_plot()
            elif self.argspec.nplot:
                self.read_table_nplot()
            elif self.argspec.plot:
                self.read_table_plot()
            elif self.argspec.run:
                self.read_table_run()
            else:
                self.read_table()
        elif op == 'Read Footer':
            if self.argspec.run:
                self.read_footer_run()
            else:
                self.read_footer()
        elif op == 'Update Header':
            if not self.argspec.paragrap_id:
                msg = str("paragrap_id required")
                raise Exception(msg)
            elif not self.argspec.update_value:
                msg = str("update_value required")
                raise Exception(msg)
            else:
                if not self.argspec.output_path:
                    self.argspec.output_path =\
                        os.path.dirname(self.argspec.docx_file)
                if self.argspec.run_id:
                    self.write_header_run()
                else:
                    self.write_header()
        elif op == 'Update Paragraphs':
            if not self.argspec.paragrap_id:
                msg = str("paragrap_id required")
                raise Exception(msg)
            elif not self.argspec.update_value:
                msg = str("update_value required")
                raise Exception(msg)
            else:
                if not self.argspec.output_path:
                    self.argspec.output_path =\
                        os.path.dirname(self.argspec.docx_file)
                if self.argspec.run_id:
                    self.write_paragraphs_run()
                else:
                    self.write_paragraphs()
        elif op == 'Update Table Rows':
            if not self.argspec.row_id:
                msg = str("Row Index required")
                raise Exception(msg)
            elif not self.argspec.cell_id:
                msg = str("cell_id required")
                raise Exception(msg)
            elif not self.argspec.update_value:
                msg = str("update_value required")
                raise Exception(msg)
            else:
                if not self.argspec.output_path:
                    self.argspec.output_path =\
                        os.path.dirname(self.argspec.docx_file)
                if self.argspec.run_id:
                    self.write_table_run()
                else:
                    self.write_table()
        elif op == 'Update Footer':
            if not self.argspec.paragrap_id:
                msg = str("paragrap_id required")
                raise Exception(msg)
            elif not self.argspec.update_value:
                msg = str("update_value required")
                raise Exception(msg)
            else:
                if not self.argspec.output_path:
                    self.argspec.output_path =\
                        os.path.dirname(self.argspec.docx_file)
                if self.argspec.run_id:
                    self.write_footer_run()
                else:
                    self.write_footer()



################################################################################
@func_log
def reg_op(mcxt, argspec):
    mcxt.logger.info('>>>starting...')
    try:
        warnings.simplefilter("ignore", ResourceWarning)
        res = Wordeditor(argspec)
        res.read_file()
        res.do(argspec.op)
        return 0
    except Exception as err:
        msg = str(err)
        mcxt.logger.error(msg)
        sys.stderr.write('%s%s' % (msg, os.linesep))
        return 1
    finally:
        sys.stdout.flush()
        mcxt.logger.info('>>>end...')


################################################################################
def _main(*args):
    with ModuleContext(
        owner='ARGOS-LABS',
        group='2',
        version='1.0',
        platform=['windows', 'darwin', 'linux'],
        output_type='text',
        display_name='Word Editor',
        icon_path=get_icon_path(__file__),
        description='Editor Module',
    ) as mcxt:
        # ##################################### for app dependent parameters

        mcxt.add_argument('docx_file', display_name='Docx File',
                          input_method='fileread',
                          help='Docx File Only')
        # ----------------------------------------------------------------------
        mcxt.add_argument('op', display_name='OP Type', help='operation types',
                          choices=['Read Header',
                                   'Read Paragraphs',
                                   'Read Table',
                                   'Read Footer',
                                   'Update Header',
                                   'Update Paragraphs',
                                   'Update Table Rows',
                                   'Update Footer'
                                   ])
        # ##################################### for app optional parameters
        mcxt.add_argument('--run', display_name='Run', action='store_true',
                          help='paragraph index has multiple formatting')
        # ----------------------------------------------------------------------
        mcxt.add_argument('--table_id', display_name='Table Index',
                          type=int,
                          default=0,
                          help='Table Index for read table row and cell location')
        # ----------------------------------------------------------------------
        mcxt.add_argument('--plot', display_name='Read Plot',
                          action='store_true',
                          help='Read Plot will plot a sample help docx')
        # ----------------------------------------------------------------------
        mcxt.add_argument('--nplot', display_name='Read Plot with New Line',
                          action='store_true',
                          help='Read Plot with New Line will plot a sample help docx')
        # ----------------------------------------------------------------------
        mcxt.add_argument('--row_id', display_name='Row Index', type=int,
                          action='append',
                          help='Row Index get from Read Table')
        # ----------------------------------------------------------------------
        mcxt.add_argument('--cell_id', display_name='Cell Index', type=int,
                          action='append',
                          help='Cell Index get from Read Table')
        # ----------------------------------------------------------------------
        mcxt.add_argument('--paragrap_id', display_name='Paragraph Index',
                          type=int,
                          action='append',
                          help='paragraph index get from reading paragraphs')
        # ----------------------------------------------------------------------
        mcxt.add_argument('--run_id', display_name='Run Index',
                          type=int,
                          action='append',
                          help='Run index get from reading paragraphs with run')
        # ----------------------------------------------------------------------
        mcxt.add_argument('--update_value', display_name='Update Value',
                          action='append',
                          help='Update paragraph, header, footer, table with this value.')
        # ----------------------------------------------------------------------
        mcxt.add_argument('--output_path', display_name='Output Path',
                          input_method='folderwrite',
                          help='An absolute folderpath to save a file')
        argspec = mcxt.parse_args(args)
        return reg_op(mcxt, argspec)


################################################################################
def main(*args):
    try:
        return _main(*args)
    except ArgsError as err:
        sys.stderr.write('Error: %s\nPlease -h to print help\n' % str(err))
    except ArgsExit as _:
        pass
