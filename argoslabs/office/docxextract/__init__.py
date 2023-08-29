#!/usr/bin/env python
# coding=utf8
"""
====================================
 :mod:`argoslabs.office.docxextract`
====================================
.. moduleauthor:: Jerry Chae <mcchae@argos-labs.com>
.. note:: ARGOS-LABS License

Description
===========
ARGOS LABS plugin module for MS Word Extract
"""
# Authors
# ===========
#
# * Jerry Chae
#
# Change Log
# --------
#
#  * [2021/04/09]
#     - 그룹에 "2-Business Apps" 넣음
#  * [2020/05/08]
#     - starting

################################################################################
import os
import sys
from alabs.common.util.vvargs import ModuleContext, func_log, \
    ArgsError, ArgsExit, get_icon_path
import csv
from io import StringIO
# noinspection PyPackageRequirements
from docx import Document
# noinspection PyPackageRequirements
from docx.document import Document as _Document
# noinspection PyPackageRequirements
from docx.oxml.text.paragraph import CT_P
# noinspection PyPackageRequirements
from docx.oxml.table import CT_Tbl
# noinspection PyPackageRequirements,PyProtectedMember
from docx.table import _Cell, Table
# noinspection PyPackageRequirements
from docx.text.paragraph import Paragraph


################################################################################
def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        # noinspection PyProtectedMember
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


################################################################################
def get_csv_table(table, is_merge=False, merge_check_col=0):
    so = StringIO()
    c = csv.writer(so, lineterminator='\n')
    p_row = None
    for i, row in enumerate(table.rows):
        row = [cell.text for cell in row.cells]
        if not is_merge:
            c.writerow(row)
            continue
        if row[merge_check_col]:
            if p_row:
                c.writerow(p_row)
            p_row = row
        else:
            if not p_row:
                p_row = row
            else:
                for j in range(len(row)):
                    if row[j]:
                        p_row[j] += '\n' + row[j]
    if p_row:
        c.writerow(p_row)
    return so.getvalue()


################################################################################
def ext_docx(document, is_merge=False, merge_check_col=0):
    so = StringIO()
    table_cnt = 0
    for block in iter_block_items(document):
        so.write('\n')
        if isinstance(block, Paragraph):
            so.write(f'{block.text}\n')
        else:
            tr = get_csv_table(document.tables[table_cnt],
                               is_merge=is_merge,
                               merge_check_col=merge_check_col)
            so.write(f'{tr}\n')
            table_cnt += 1
    return so.getvalue()


################################################################################
@func_log
def docx_extract(mcxt, argspec):
    mcxt.logger.info('>>>starting...')
    try:
        if not os.path.exists(argspec.docx):
            raise IOError(f'Cannot read docx file "{argspec.docx}"')
        _, ext = os.path.splitext(argspec.docx)
        if ext.lower() != '.docx':
            raise RuntimeError(f'File extension must ".docx" but "{ext}"')
        document = Document(argspec.docx)
        if argspec.table_only:
            if not (0 <= argspec.table_index < len(document.tables) or
                    -len(document.tables) <= argspec.table_index < 0):
                raise ReferenceError(f'Number of tables: {len(document.tables)} '
                                     f'but invalid index {argspec.table_index}')
            r = get_csv_table(document.tables[argspec.table_index],
                              is_merge=argspec.cell_merge,
                              merge_check_col=argspec.merge_check_col)
        else:
            r = ext_docx(document,
                         is_merge=argspec.cell_merge,
                         merge_check_col=argspec.merge_check_col)
        print(r, end='')
        return 0
    except Exception as err:
        msg = str(err)
        mcxt.logger.error(msg)
        sys.stderr.write('%s%s' % (msg, os.linesep))
        return 1
    finally:
        sys.stdout.flush()
        mcxt.logger.info('>>>end...')


################################################################################
def _main(*args):
    with ModuleContext(
        owner='ARGOS-LABS',
        group='2',  # Business Apps
        version='1.0',
        platform=['windows', 'darwin', 'linux'],
        output_type='text',
        display_name='MS Word Extract',
        icon_path=get_icon_path(__file__),
        description='Extract text form MS Word (docx)',
    ) as mcxt:
        # ##################################### for app dependent parameters
        mcxt.add_argument('docx', action='store',
                          display_name='DOCX File',
                          input_method='fileread',
                          help='docx MS word file to extract text')
        # ######################################## for app dependent options
        mcxt.add_argument('--table-only', action='store_true',
                          display_name='Table Only',
                          help='If set this flag extract table only')
        mcxt.add_argument('--table-index', type=int,
                          default=0,
                          display_name='Table Index',
                          help='Table index to extract table, default is [[0]], first table')
        mcxt.add_argument('--cell-merge', action='store_true',
                          display_name='Cell Merge',
                          help='If set this flag then merge contents for empty Index Column')
        mcxt.add_argument('--merge-check-col', type=int,
                          default=0,
                          display_name='Merge Check Col',
                          help='If Cell Merge is checked and this column is empty then merged')

        argspec = mcxt.parse_args(args)
        return docx_extract(mcxt, argspec)


################################################################################
def main(*args):
    try:
        return _main(*args)
    except ArgsError as err:
        sys.stderr.write('Error: %s\nPlease -h to print help\n' % str(err))
    except ArgsExit as _:
        pass
