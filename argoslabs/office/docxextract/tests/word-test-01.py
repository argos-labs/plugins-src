import sys
import csv
from docx import Document
from io import StringIO


################################################################################
def get_csv_table(table, is_merge=False, index_col=0):
    so = StringIO()
    c = csv.writer(so, lineterminator='\n')
    p_row = None
    for i, row in enumerate(table.rows):
        row = [cell.text for cell in row.cells]
        if not is_merge:
            c.writerow(row)
            continue
        if row[index_col]:
            if p_row:
                c.writerow(p_row)
                p_row = row
            else:
                p_row = row
        else:
            if not p_row:
                p_row = row
            else:
                for i in range(len(row)):
                    if row[i]:
                        p_row[i] += '_n' + row[i]
    if p_row:
        c.writerow(p_row)
    return so.getvalue()

# document = Document('foo-01.docx')
# table = document.tables[-1]
# data = []
# keys = {}

# c = csv.writer(sys.stdout, lineterminator='\n')
# for i, row in enumerate(table.rows):
#     text = [cell.text for cell in row.cells]
#     c.writerow(text)

# print(data)


document = Document('foo-01.docx')
for table in document.tables:
    print('%s' % ('=' * 80))
    print(get_csv_table(table, is_merge=True))
